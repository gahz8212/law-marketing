import io
import re
import zipfile
from typing import Optional
from pypdf import PdfReader
import docx
from google import genai
from google.genai import types
from app.core.config import LLM_API_KEY, DEFAULT_GEMINI_MODEL

class DocumentParser:
    """
    변호사가 전달한 모든 형태의 자료(PDF, HWP, HWPX, Word DOCX, TXT, 이미지 PNG/JPG)에서
    순수 텍스트를 고속으로 추출하는 올인원 멀티포맷 파서
    """

    IMAGE_EXTENSIONS = ["png", "jpg", "jpeg", "webp", "bmp"]

    @classmethod
    def extract_text(cls, file_bytes: bytes, filename: str) -> str:
        ext = filename.lower().split(".")[-1] if "." in filename else ""

        if ext == "pdf":
            return cls._extract_from_pdf(file_bytes)
        elif ext in ["docx", "doc"]:
            return cls._extract_from_docx(file_bytes)
        elif ext == "hwpx":
            return cls._extract_from_hwpx(file_bytes)
        elif ext == "hwp":
            return cls._extract_from_hwp(file_bytes)
        elif ext in cls.IMAGE_EXTENSIONS:
            return cls._extract_from_image(file_bytes, ext)
        elif ext in ["txt", "md", "csv"]:
            return cls._extract_from_text(file_bytes)
        else:
            # 기본 텍스트 디코딩
            return cls._extract_from_text(file_bytes)

    @staticmethod
    def _extract_from_pdf(file_bytes: bytes) -> str:
        """PDF 파일에서 텍스트 추출 (전자소송 정본/등본 지원)"""
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                txt = page.extract_text() or ""
                if txt.strip():
                    pages_text.append(txt.strip())
            return "\n\n".join(pages_text)
        except Exception as e:
            return f"PDF 텍스트 추출 실패: {str(e)}"

    @staticmethod
    def _extract_from_docx(file_bytes: bytes) -> str:
        """MS Word (.docx) 파일에서 텍스트 추출"""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            # 표(Table) 내부 텍스트도 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            return "\n".join(full_text)
        except Exception as e:
            return f"워드(DOCX) 텍스트 추출 실패: {str(e)}"

    @staticmethod
    def _extract_from_hwpx(file_bytes: bytes) -> str:
        """한글 HWPX(ZIP XML 기반) 파일에서 텍스트 추출"""
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                section_files = [f for f in z.namelist() if f.startswith("Contents/section") and f.endswith(".xml")]
                if not section_files:
                    section_files = [f for f in z.namelist() if f.endswith(".xml")]

                extracted = []
                for s_file in section_files:
                    xml_content = z.read(s_file).decode("utf-8", errors="ignore")
                    clean = re.sub(r"<[^>]+>", " ", xml_content)
                    clean = re.sub(r"\s+", " ", clean).strip()
                    if clean:
                        extracted.append(clean)
                return "\n\n".join(extracted)
        except Exception as e:
            return f"HWPX 텍스트 추출 실패: {str(e)}"

    @staticmethod
    def _extract_from_hwp(file_bytes: bytes) -> str:
        """바이너리 HWP 5.0 파일에서 본문 텍스트 패턴 추출"""
        try:
            decoded = file_bytes.decode("utf-16le", errors="ignore")
            korean_blocks = re.findall(r"[가-힣0-9a-zA-Z\s.,·()\-]{4,}", decoded)
            clean = " ".join(korean_blocks)
            clean = re.sub(r"\s+", " ", clean).strip()
            if len(clean) > 20:
                return clean
            fallback = file_bytes.decode("cp949", errors="ignore")
            korean_blocks_fb = re.findall(r"[가-힣0-9a-zA-Z\s.,·()\-]{4,}", fallback)
            return " ".join(korean_blocks_fb).strip()
        except Exception as e:
            return f"HWP 텍스트 추출 실패: {str(e)}"

    @classmethod
    def _extract_from_image(cls, file_bytes: bytes, ext: str) -> str:
        """
        변호사가 스마트폰으로 찍은 판결문 사진이나 화면 캡처 이미지(PNG/JPG)에서
        Gemini 멀티모달 Vision AI를 활용하여 텍스트를 정확하게 판독(OCR) 추출
        """
        if not LLM_API_KEY:
            return "이미지 판독을 위한 LLM_API_KEY가 설정되지 않았습니다."

        mime_type = "image/png" if ext == "png" else "image/jpeg"
        if ext == "webp":
            mime_type = "image/webp"

        try:
            client = genai.Client(api_key=LLM_API_KEY)
            prompt = (
                "이 이미지는 변호사가 전달한 법원 판결문, 소장, 또는 승소 관련 서류입니다.\n"
                "이미지에 적힌 모든 텍스트(주문, 이유, 사실관계, 사건번호, 당사자 주장 등)를 "
                "누락 없이 한국어로 빠짐없이 읽어서 정확하게 텍스트로 옮겨 적어주세요. "
                "설명이나 부연 없이 이미지 본문 텍스트만 그대로 출력하세요."
            )
            image_part = types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
            resp = client.models.generate_content(
                model=DEFAULT_GEMINI_MODEL,
                contents=[image_part, prompt]
            )
            return resp.text.strip()
        except Exception as e:
            return f"이미지 판독(OCR) 실패: {str(e)}"

    @staticmethod
    def _extract_from_text(file_bytes: bytes) -> str:
        """일반 텍스트 파일(UTF-8 / CP949 자동 감지) 디코딩"""
        for encoding in ["utf-8", "cp949", "euc-kr"]:
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="ignore")
